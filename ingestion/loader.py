from pathlib import Path
import re

import pandas as pd

from pypdf import PdfReader
from docx import Document as DocxDocument
from tempfile import NamedTemporaryFile

from langchain_core.documents import Document

# =========================
# Extensions
# =========================

TEXT_EXTENSIONS = {
    ".txt",
    ".md",
    ".csv",
    ".json",
    ".py",
    ".html",
    ".xml",
    ".log",
}

EXCEL_EXTENSIONS = {
    ".xlsx",
    ".xls",
}

# =========================
# Safe Filename
# =========================

def safe_filename(filename: str) -> str:

    name = Path(filename).name

    return re.sub(
        r"[^A-Za-z0-9._ -]",
        "_",
        name,
    )

# =========================
# Save Uploaded Files
# =========================

def save_uploaded_files(uploaded_files):

    if not uploaded_files:
        return

    upload_dir = Path(
        "data/uploaded_docs"
    )

    upload_dir.mkdir(
        parents=True,
        exist_ok=True,
    )

    for uploaded_file in uploaded_files:

        destination = (
            upload_dir
            / safe_filename(
                uploaded_file.name
            )
        )

        destination.write_bytes(
            uploaded_file.getvalue()
        )

# =========================
# Load Text File
# =========================

def load_text_file(
    file,
    source_name,
    origin,
):

    if isinstance(file, (str, Path)):

        raw_text = Path(file).read_text(
            encoding="utf-8",
            errors="ignore",
        )

    else:

        raw_text = file.read().decode(
            "utf-8",
            errors="ignore",
        )

    if not raw_text.strip():
        return []

    return [
        Document(
            page_content=raw_text,
            metadata={
                "source": source_name,
                "origin": origin,
            },
        )
    ]

# =========================
# Load PDF
# =========================

def load_pdf(
    file,
    source_name,
    origin,
):

    reader = PdfReader(file)

    documents = []

    for page_number, page in enumerate(
        reader.pages,
        start=1,
    ):

        text = page.extract_text() or ""

        if not text.strip():
            continue

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": source_name,
                    "page": page_number,
                    "origin": origin,
                },
            )
        )

    return documents

# =========================
# Load Excel
# =========================

def load_excel(
    file,
    source_name,
    origin,
):

    try:

        sheets = pd.read_excel(
            file,
            sheet_name=None,
        )

    except Exception:

        return []

    documents = []

    for sheet_name, dataframe in sheets.items():

        dataframe = dataframe.dropna(
            how="all"
        ).dropna(
            axis=1,
            how="all",
        )

        if dataframe.empty:
            continue

        text = dataframe.to_string()

        documents.append(
            Document(
                page_content=text,
                metadata={
                    "source": source_name,
                    "sheet": sheet_name,
                    "origin": origin,
                },
            )
        )

    return documents


def load_docx(
    file,
    source_name,
    origin,
):

    try:

        if isinstance(file, (str, Path)):

            doc = DocxDocument(file)

        else:

            with NamedTemporaryFile(
                delete=False,
                suffix=".docx",
            ) as temp_file:

                temp_file.write(
                    file.read()
                )

                temp_path = temp_file.name

            doc = DocxDocument(
                temp_path
            )

        text = "\n".join(
            paragraph.text
            for paragraph in doc.paragraphs
        )

        if not text.strip():
            return []

        return [
            Document(
                page_content=text,
                metadata={
                    "source": source_name,
                    "origin": origin,
                },
            )
        ]

    except Exception:

        return []

# =========================
# Load File
# =========================

def load_file(
    file,
    source_name,
    origin,
):

    suffix = Path(
        source_name
    ).suffix.lower()

    if suffix == ".pdf":

        return load_pdf(
            file,
            source_name,
            origin,
        )

    if suffix in EXCEL_EXTENSIONS:

        return load_excel(
            file,
            source_name,
            origin,
        )

    if suffix == ".docx":
       return load_docx(
        file,
        source_name,
        origin,
    )

    return load_text_file(
        file,
        source_name,
        origin,
    )
    

# =========================
# Load Documents Folder
# =========================

def load_documents_folder(
    folder,
    origin,
):

    documents = []

    if not folder.exists():

        return documents

    for path in sorted(
        folder.rglob("*")
    ):

        if not path.is_file():
            continue

        documents.extend(
            load_file(
                path,
                source_name=str(
                    path.relative_to(
                        folder
                    )
                ),
                origin=origin,
            )
        )

    return documents